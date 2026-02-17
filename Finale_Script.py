"""
Technical Data Analysis & ETL Pipeline – GBP_DataSource

The Script is designed to:
1. Assess data quality (Completeness, Inconsistencies, Redundancies, Duplicates).
2. Perform Inconsistency Mapping (One-to-Many relationship violations).
3. Derive KPIs and Business Attributes.
4. Clean and Standardize the dataset for downstream analysis.

Author: Innocent Amos Mchechesi
Date: 2026
"""

import os
import io
import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Pt
from typing import Union, List

# ==============================
# 1. Configuration & Constants
# ==============================
RAW_DATA_PATH = "GBP_DataSource_Pr.csv"
CLEAN_DATA_PATH = "GBP_DataSource_clean.csv"
LOG_PATH = "GBP_DataSource_data_quality_log.txt"
REPORT_PATH = "GBP_DataSource_Analysis_Report.docx"

# Business Logic Configuration
PRIMARY_KEY_COLS = ["Order ID", "Row ID"] 
DATE_COLS = ["Order Date", "Ship Date"]
NUMERIC_COLS = ["Sales", "Quantity", "Discount", "Profit"]

# ==============================
# 2. Professional Logging Engine
# ==============================
class AnalysisLogger:
    def __init__(self, txt_path, doc_path):
        self.txt_path = txt_path
        self.doc_path = doc_path
        self.doc = Document()
        self.doc.add_heading('Technical Data Analysis Report', 0)
        
        # Initialize text file
        with open(self.txt_path, "w", encoding="utf-8") as f:
            f.write(f"DATA ANALYSIS LOG | Generated: {datetime.now()}\n")
            f.write("="*80 + "\n")

    def log_section(self, title, level=1):
        """Adds a section header to console, txt, and docx."""
        print(f"\n--- {title.upper()} ---")
        with open(self.txt_path, "a", encoding="utf-8") as f:
            f.write(f"\n\nSECTION: {title}\n" + "-"*40 + "\n")
        
        self.doc.add_heading(title, level=level)

    def log_message(self, message):
        """Logs a simple string message."""
        print(message)
        with open(self.txt_path, "a", encoding="utf-8") as f:
            f.write(f"{datetime.now().strftime('%H:%M:%S')} | {message}\n")
        self.doc.add_paragraph(message)

    def log_table(self, df, title=None, max_rows=15):
        """Converts a DataFrame to a formatted table in the Word Doc."""
        if title:
            self.log_message(title)
        
        if df.empty:
            self.doc.add_paragraph("No issues/data found for this metric.")
            return

        # Limit size for report readability
        display_df = df.head(max_rows).reset_index() if max_rows else df.reset_index()
        
        table = self.doc.add_table(rows=1, cols=len(display_df.columns))
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(display_df.columns):
            hdr_cells[i].text = str(col_name)
            paragraph = hdr_cells[i].paragraphs[0]
            run = paragraph.runs[0]
            run.bold = True

        # Data Rows
        for _, row in display_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                if isinstance(val, float):
                    row_cells[i].text = f"{val:.2f}"
                else:
                    row_cells[i].text = str(val)
        
        if len(df) > max_rows:
            self.doc.add_paragraph(f"... [Truncated: {len(df) - max_rows} more rows]")
        self.doc.add_paragraph() # Spacer

    def save_report(self):
        self.doc.save(self.doc_path)

# ==============================
# 3. Core Analysis Logic
# ==============================
def detect_inconsistencies(df: pd.DataFrame, col1: str, col2: Union[str, List[str]]):
    """
    Identifies cases where one value in col1 maps to multiple different values in col2.
    Example: One 'Postal Code' associated with two different 'Cities'.
    """
    df_temp = df.copy()
    if isinstance(col2, str): col2 = [col2]
    
    check_cols = [col1] + col2
    # Drop NAs for this check to avoid false positives
    df_temp = df_temp.dropna(subset=check_cols)
    
    # Filter groups where col1 has more than 1 unique value in any col2
    inconsistent = df_temp.groupby(col1).filter(lambda x: x[col2].nunique().gt(1).any())
    return inconsistent[check_cols].drop_duplicates().sort_values(by=col1)

def run_pipeline():
    logger = AnalysisLogger(LOG_PATH, REPORT_PATH)

    # 1. LOAD DATA
    if not os.path.exists(RAW_DATA_PATH):
        logger.log_message(f"CRITICAL: {RAW_DATA_PATH} not found.")
        return

    df = pd.read_csv(RAW_DATA_PATH, delimiter=';', decimal=',', encoding='utf-8-sig')
    logger.log_section("1. Data Quality Assessment")
    logger.log_message(f"Initial Shape: {df.shape[0]} rows, {df.shape[1]} columns.")

    # 2. ASSESSMENT: Completeness & Duplicates
    missing = df.isna().sum().to_frame(name='Missing Count')
    logger.log_table(missing[missing['Missing Count'] > 0], "Missing Values Summary")

    full_dups = df.duplicated().sum()
    logger.log_message(f"Fully Duplicated Rows: {full_dups}")

    # 3. ASSESSMENT: Inconsistencies (Relational)
    logger.log_section("2. Logical Inconsistency Check", level=2)
    
    # Check: Postal Code -> City
    if {'Postal Code', 'City'}.issubset(df.columns):
        inc_postal = detect_inconsistencies(df, 'Postal Code', 'City')
        logger.log_table(inc_postal, "Inconsistent Mapping: One Postal Code to Many Cities")

    # Check: Product ID -> Product Name
    if {'Product ID', 'Product Name'}.issubset(df.columns):
        inc_prod = detect_inconsistencies(df, 'Product ID', 'Product Name')
        logger.log_table(inc_prod, "Inconsistent Mapping: One Product ID to Many Names")

    # 4. KPI & ATTRIBUTE DEFINITION
    logger.log_section("3. Calculations & KPI Enrichment")
    
    # Example KPI: Profit Margin
    if {'Profit', 'Sales'}.issubset(df.columns):
        df['Profit_Margin'] = (df['Profit'] / df['Sales']).replace([np.inf, -np.inf], 0)
        logger.log_message("Added KPI: Profit_Margin (Profit / Sales)")

    # Derived Entity
    if {'Product ID', 'Product Name'}.issubset(df.columns):
        df['Product_Full_Name'] = df['Product ID'].astype(str) + " | " + df['Product Name'].astype(str)
        logger.log_message("Added Attribute: Product_Full_Name")

    # 5. DATA CLEANING
    logger.log_section("4. Data Cleaning & Transformation Log")
    
    # Standardization
    df.columns = df.columns.str.strip().str.replace(' ', '_').str.lower()
    logger.log_message("Standardized column names to lowercase snake_case.")

    # Drop Duplicates
    initial_count = len(df)
    df = df.drop_duplicates()
    logger.log_message(f"Dropped {initial_count - len(df)} duplicate rows.")

    # Handle Dates
    for col in [c.lower().replace(' ', '_') for c in DATE_COLS]:
        if col in df.columns:
            df[col] = pd.to_datetime(df[col], errors='coerce')
            logger.log_message(f"Converted '{col}' to datetime.")

    # Handle Missing Values (Imputation)
    num_cols_actual = df.select_dtypes(include=[np.number]).columns
    for col in num_cols_actual:
        if df[col].isna().any():
            median_val = df[col].median()
            df[col] = df[col].fillna(median_val)
            logger.log_message(f"Imputed missing values in '{col}' with median: {median_val}")

    # 6. FINALIZATION
    logger.log_section("5. Final Dataset Summary")
    logger.log_message(f"Final Row Count: {len(df)}")
    df.to_csv(CLEAN_DATA_PATH, index=False)
    logger.log_message(f"Cleaned dataset exported to: {CLEAN_DATA_PATH}")

    logger.save_report()
    print(f"\nProcess Complete. See '{REPORT_PATH}' for the full analysis.")

if __name__ == "__main__":
    run_pipeline()